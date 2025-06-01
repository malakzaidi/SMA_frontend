from flask import Blueprint, request, jsonify, send_file, session
import os
import sys
import json
import tempfile
import subprocess
from werkzeug.utils import secure_filename
import speech_recognition as sr
from io import BytesIO
import markdown
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
from datetime import datetime

user_bp = Blueprint('user', __name__)

# In-memory user database (would be replaced by a real database in production)
users = {}
reports = {}

@user_bp.route('/auth/register', methods=['POST'])
def register():
    """Register a new user."""
    data = request.json
    if not data or 'email' not in data or 'password' not in data or 'name' not in data:
        return jsonify({'error': 'Missing required fields'}), 400
    
    email = data['email']
    password = data['password']
    name = data['name']
    
    if email in users:
        return jsonify({'error': 'Email already registered'}), 409
    
    user_id = str(uuid.uuid4())
    users[email] = {
        'id': user_id,
        'email': email,
        'password': password,  # In a real app, this would be hashed
        'name': name,
        'created_at': datetime.now().isoformat()
    }
    
    reports[user_id] = []
    
    return jsonify({
        'id': user_id,
        'email': email,
        'name': name
    }), 201

@user_bp.route('/auth/login', methods=['POST'])
def login():
    """Login a user."""
    data = request.json
    if not data or 'email' not in data or 'password' not in data:
        return jsonify({'error': 'Missing email or password'}), 400
    
    email = data['email']
    password = data['password']
    
    if email not in users or users[email]['password'] != password:
        return jsonify({'error': 'Invalid email or password'}), 401
    
    user = users[email]
    session['user_id'] = user['id']
    
    return jsonify({
        'id': user['id'],
        'email': user['email'],
        'name': user['name']
    }), 200

@user_bp.route('/auth/logout', methods=['POST'])
def logout():
    """Logout a user."""
    session.pop('user_id', None)
    return jsonify({'message': 'Logged out successfully'}), 200

@user_bp.route('/reports', methods=['GET'])
def get_reports():
    """Get all reports for the current user."""
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'error': 'Not authenticated'}), 401
    
    user_reports = reports.get(user_id, [])
    return jsonify(user_reports), 200

@user_bp.route('/transcribe', methods=['POST'])
def transcribe_audio():
    """Transcribe audio to text using speech recognition."""
    if 'audio' not in request.files:
        return jsonify({'error': 'No audio file provided'}), 400
    
    audio_file = request.files['audio']
    
    # Save the audio file temporarily
    temp_dir = tempfile.mkdtemp()
    temp_audio_path = os.path.join(temp_dir, secure_filename('audio.webm'))
    audio_file.save(temp_audio_path)
    
    try:
        # Convert webm to wav using ffmpeg
        temp_wav_path = os.path.join(temp_dir, 'audio.wav')
        os.system(f'ffmpeg -i {temp_audio_path} {temp_wav_path} -y')
        
        # Use speech recognition to transcribe
        recognizer = sr.Recognizer()
        with sr.AudioFile(temp_wav_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data, language='fr-FR')
        
        # Clean up temporary files
        os.remove(temp_audio_path)
        os.remove(temp_wav_path)
        os.rmdir(temp_dir)
        
        return jsonify({'text': text}), 200
    
    except Exception as e:
        # Clean up in case of error
        if os.path.exists(temp_audio_path):
            os.remove(temp_audio_path)
        if os.path.exists(os.path.join(temp_dir, 'audio.wav')):
            os.remove(os.path.join(temp_dir, 'audio.wav'))
        if os.path.exists(temp_dir):
            os.rmdir(temp_dir)
        
        print(f"Error in transcription: {e}")
        return jsonify({'error': str(e)}), 500

@user_bp.route('/generate', methods=['POST'])
def generate_report():
    """Generate a radiology report based on the provided prompt."""
    data = request.json
    if not data or 'prompt' not in data:
        return jsonify({'error': 'No prompt provided'}), 400
    
    prompt = data['prompt']
    
    try:
        # First, determine the exam type from the prompt
        exam_type = determine_exam_type(prompt)
        
        # Generate a report based on the exam type
        if "cérébrale" in prompt.lower() or "cerveau" in prompt.lower() or "tête" in prompt.lower():
            report_content = generate_cerebral_mri_report(prompt)
        elif "genou" in prompt.lower():
            report_content = generate_knee_mri_report(prompt)
        elif "rachis" in prompt.lower() or "lombaire" in prompt.lower() or "colonne" in prompt.lower():
            report_content = generate_spine_mri_report(prompt)
        else:
            # Generic report if no specific type is detected
            report_content = generate_generic_mri_report(prompt)
        
        # Save report to history if user is authenticated
        user_id = session.get('user_id')
        if user_id:
            report_id = str(uuid.uuid4())
            report_data = {
                'id': report_id,
                'prompt': prompt,
                'report': report_content,
                'exam_type': exam_type,
                'created_at': datetime.now().isoformat()
            }
            
            if user_id not in reports:
                reports[user_id] = []
            
            reports[user_id].append(report_data)
        
        return jsonify({
            'report': report_content,
            'exam_type': exam_type
        }), 200
    
    except Exception as e:
        print(f"Error generating report: {e}")
        return jsonify({'error': str(e)}), 500

@user_bp.route('/download', methods=['POST'])
def download_report():
    """Convert the report to DOCX and send it for download."""
    data = request.json
    if not data or 'report' not in data:
        return jsonify({'error': 'No report content provided'}), 400
    
    report_content = data['report']
    
    try:
        # Create a new DOCX document
        doc = docx.Document()
        
        # Set document properties
        doc.core_properties.title = "Rapport Radiologique"
        doc.core_properties.author = "Système Multi-Agents"
        
        # Parse markdown content
        lines = report_content.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Handle headings
            if line.startswith('# '):
                # Main title
                title = line[2:].strip()
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(63, 81, 181)  # #3f51b5
                run.font.size = Pt(18)
                doc.add_paragraph()
            elif line.startswith('## '):
                # Section heading
                current_section = line[3:].strip()
                heading = doc.add_heading(current_section, 1)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(48, 63, 159)  # #303f9f
                run.font.size = Pt(16)
            elif line.startswith('### '):
                # Subsection heading
                subsection = line[4:].strip()
                heading = doc.add_heading(subsection, 2)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(48, 63, 159)  # #303f9f
                run.font.size = Pt(14)
            elif line.startswith('#### '):
                # Sub-subsection heading
                subsubsection = line[5:].strip()
                heading = doc.add_heading(subsubsection, 3)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(48, 63, 159)  # #303f9f
                run.font.size = Pt(13)
            else:
                # Regular paragraph
                para = doc.add_paragraph(line)
                para.style = 'Normal'
                for run in para.runs:
                    run.font.size = Pt(11)
        
        # Save the document to a BytesIO object
        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        return send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='rapport_radiologique.docx'
        )
    
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        return jsonify({'error': str(e)}), 500

def determine_exam_type(prompt):
    """Determine the exam type from the prompt."""
    if "cérébrale" in prompt.lower() or "cerveau" in prompt.lower() or "tête" in prompt.lower():
        return "IRM cérébrale"
    elif "genou" in prompt.lower():
        return "IRM du genou"
    elif "rachis" in prompt.lower() or "lombaire" in prompt.lower() or "colonne" in prompt.lower():
        return "IRM rachidienne"
    elif "épaule" in prompt.lower():
        return "IRM de l'épaule"
    elif "hanche" in prompt.lower():
        return "IRM de la hanche"
    elif "foie" in prompt.lower() or "hépati" in prompt.lower():
        return "IRM hépatique"
    elif "pelvis" in prompt.lower() or "pelvienne" in prompt.lower():
        return "IRM pelvienne"
    else:
        return "IRM générique"

def generate_cerebral_mri_report(prompt):
    """Generate a cerebral MRI report."""
    return f"""# Rapport IRM Cérébrale

## Indication
{prompt}

## Technique
IRM cérébrale réalisée sur appareil 1.5 Tesla avec séquences T1, T2, FLAIR, diffusion et T1 après injection de gadolinium.

## Résultats
- **Parenchyme cérébral**: Absence d'anomalie de signal parenchymateuse. Pas de lésion ischémique récente visible en diffusion.
- **Système ventriculaire**: Taille et morphologie normales.
- **Espaces sous-arachnoïdiens**: Non élargis.
- **Structures médianes**: En place.
- **Fosse postérieure**: Absence d'anomalie cérébelleuse ou du tronc cérébral.
- **Sinus de la face**: Sans particularité.

## Conclusion
IRM cérébrale sans anomalie significative. Absence de signe d'accident vasculaire cérébral récent.
"""

def generate_knee_mri_report(prompt):
    """Generate a knee MRI report."""
    return f"""# Rapport IRM du Genou

## Indication
{prompt}

## Technique
IRM du genou réalisée sur appareil 1.5 Tesla avec séquences DP FS dans les 3 plans, T1 sagittale.

## Résultats
- **Cartilage**: Cartilage fémoro-tibial et fémoro-patellaire d'épaisseur normale.
- **Ménisques**: Absence de lésion méniscale interne ou externe.
- **Ligaments croisés**: LCA et LCP d'aspect normal.
- **Ligaments collatéraux**: Ligaments collatéraux médial et latéral sans anomalie.
- **Tendons**: Tendon rotulien d'aspect normal.
- **Épanchement**: Absence d'épanchement articulaire significatif.

## Conclusion
IRM du genou sans anomalie significative.
"""

def generate_spine_mri_report(prompt):
    """Generate a spine MRI report."""
    return f"""# Rapport IRM Rachidienne

## Indication
{prompt}

## Technique
IRM du rachis lombaire réalisée sur appareil 1.5 Tesla avec séquences T2 sagittale, T1 sagittale, STIR sagittale, T2 axiale centrée sur L3-L5.

## Résultats
- **Alignement vertébral**: Conservation des courbures physiologiques.
- **Corps vertébraux**: Hauteur conservée. Absence de tassement vertébral.
- **Disques intervertébraux**: 
  * L3-L4: Discopathie dégénérative modérée sans hernie discale.
  * L4-L5: Protrusion discale médiane sans compression radiculaire.
  * L5-S1: Discopathie dégénérative avec pincement discal modéré.
- **Canal rachidien**: Absence de sténose canalaire significative.
- **Foramens**: Absence de sténose foraminale significative.
- **Articulaires postérieures**: Arthrose facettaire modérée en L4-L5 et L5-S1.

## Conclusion
Discopathie dégénérative lombaire modérée prédominant en L4-L5 et L5-S1, sans compression radiculaire significative.
"""

def generate_generic_mri_report(prompt):
    """Generate a generic MRI report."""
    return f"""# Rapport IRM

## Indication
{prompt}

## Technique
Examen réalisé selon le protocole standard.

## Résultats
L'examen ne montre pas d'anomalie significative.

## Conclusion
IRM sans particularité dans les limites de l'examen réalisé.
"""
